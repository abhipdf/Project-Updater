"""
Word document generator for Project Update Studio.
Generates comprehensive project documentation.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
from utils import COLORS, format_date, format_datetime, safe_json_parse, get_string
from typing import List, Dict
import os


def hex_to_rgb(hex_color: str) -> tuple:
    """Convert hex color to RGB tuple."""
    hex_color = hex_color.lstrip("#")
    return tuple(int(hex_color[i : i + 2], 16) for i in (0, 2, 4))


def generate_project_documentation(
    project: Dict,
    team_members: List[Dict],
    all_updates: List[Dict],
    ai_closure_summary: str,
    language: str = "en",
    filepath: str = None,
) -> str:
    """Generate comprehensive Word document for a project.
    
    Args:
        project: Project data dict
        team_members: List of team member dicts
        all_updates: List of weekly update dicts (sorted by week)
        ai_closure_summary: AI-generated project closure summary
        language: Language code (en/de)
        filepath: Output file path
        
    Returns:
        Path to generated file
    """
    if filepath is None:
        now = datetime.now().strftime("%Y%m%d_%H%M%S")
        project_name = project["name"].replace(" ", "_")
        exports_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "exports")
        os.makedirs(exports_dir, exist_ok=True)
        filepath = os.path.join(exports_dir, f"{project_name}_{now}.docx")

    doc = Document()
    
    # Set up styles
    style = doc.styles["Normal"]
    style.font.size = Pt(11)
    style.font.name = "Calibri"

    # Color scheme
    primary_rgb = hex_to_rgb(COLORS["primary_blue"])
    dark_rgb = hex_to_rgb(COLORS["dark_blue"])

    # === COVER PAGE ===
    cover_title = doc.add_paragraph()
    cover_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cover_title.add_run(project["name"])
    run.font.size = Pt(48)
    run.font.bold = True
    run.font.color.rgb = RGBColor(*dark_rgb)

    doc.add_paragraph()  # Spacing

    # Project period
    if all_updates:
        first_date = all_updates[-1].get("created_at", "")
        last_date = all_updates[0].get("created_at", "")
        date_range = f"{format_date(first_date, language)} – {format_date(last_date, language)}"
    else:
        date_range = format_date(datetime.now().isoformat(), language)

    period_para = doc.add_paragraph(date_range)
    period_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    period_para.runs[0].font.size = Pt(14)
    period_para.runs[0].font.italic = True

    doc.add_paragraph()  # Spacing

    # Team members
    if team_members:
        team_heading = doc.add_paragraph(get_string("team", language))
        team_heading.runs[0].font.size = Pt(12)
        team_heading.runs[0].font.bold = True
        team_heading.runs[0].font.color.rgb = RGBColor(*primary_rgb)

        for member in team_members:
            member_str = f"{member['name']}"
            if member.get("role"):
                member_str += f" – {member['role']}"
            doc.add_paragraph(member_str, style="List Bullet")

    doc.add_paragraph()  # Spacing
    doc.add_paragraph()  # Spacing

    # Final status
    final_rag = project.get("rag_status", "green")
    rag_label_key = f"rag_{final_rag}"
    rag_label = get_string(rag_label_key, language)

    status_para = doc.add_paragraph()
    status_para.add_run(f"{get_string('rag_green', language) if final_rag == 'green' else get_string('rag_amber', language) if final_rag == 'amber' else get_string('rag_red', language)}: ").font.bold = True
    status_para.add_run(final_rag.upper())

    # Generate date
    gen_date = format_datetime(datetime.now().isoformat(), language)
    doc.add_paragraph(f"{get_string('loading', language)}: {gen_date}", style="Normal")

    # Page break
    doc.add_page_break()

    # === SECTION 1: PROJECT OVERVIEW ===
    heading1 = doc.add_heading(get_string("project_goal", language), level=1)
    heading1.runs[0].font.color.rgb = RGBColor(*dark_rgb)

    if project.get("description"):
        doc.add_heading(get_string("project_description", language), level=2).runs[0].font.color.rgb = RGBColor(*primary_rgb)
        doc.add_paragraph(project["description"])

    if project.get("goal"):
        doc.add_heading(get_string("project_goal", language), level=2).runs[0].font.color.rgb = RGBColor(*primary_rgb)
        doc.add_paragraph(project["goal"])

    if project.get("background"):
        doc.add_heading(get_string("project_background", language), level=2).runs[0].font.color.rgb = RGBColor(*primary_rgb)
        doc.add_paragraph(project["background"])

    doc.add_page_break()

    # === SECTION 2: WEEKLY UPDATES LOG ===
    heading2 = doc.add_heading(get_string("weekly_summary", language), level=1)
    heading2.runs[0].font.color.rgb = RGBColor(*dark_rgb)

    for update in all_updates:
        # Week heading
        week_heading = doc.add_heading(
            f"{update.get('week_label', 'Week')} – {format_date(update.get('created_at', ''), language)}",
            level=2
        )
        week_heading.runs[0].font.color.rgb = RGBColor(*primary_rgb)

        # RAG status
        rag_status = update.get("rag_status", "green")
        rag_para = doc.add_paragraph()
        rag_para.add_run(get_string("rag_green", language) + ": ").bold = True
        rag_para.add_run(rag_status.upper())

        # Summary
        if update.get("ai_summary"):
            doc.add_paragraph(update["ai_summary"], style="Normal")

        # Tasks completed
        tasks = safe_json_parse(update.get("tasks_completed", "[]"), [])
        if tasks:
            doc.add_heading(get_string("tasks_completed", language), level=3)
            for task in tasks:
                if isinstance(task, dict):
                    task_text = f"{task.get('task', '')} → {task.get('result','')}"
                    if task.get("owner"):
                        task_text += f" ({task['owner']})"
                else:
                    task_text = str(task)
                doc.add_paragraph(task_text, style="List Bullet")

        # Next tasks
        next_tasks = safe_json_parse(update.get("next_tasks", "[]"), [])
        if next_tasks:
            doc.add_heading(get_string("next_steps", language), level=3)
            for task in next_tasks:
                if isinstance(task, dict):
                    task_text = f"{task.get('task', '')}"
                    if task.get("owner"):
                        task_text += f" ({task['owner']})"
                    if task.get("due_date"):
                        task_text += f" – {task['due_date']}"
                else:
                    task_text = str(task)
                doc.add_paragraph(task_text, style="List Bullet")

        # Risks & blockers
        risks = safe_json_parse(update.get("risks_blockers", "[]"), [])
        if risks:
            doc.add_heading(get_string("risks_blockers", language), level=3)
            for risk in risks:
                if isinstance(risk, dict):
                    risk_text = f"{risk.get('issue', '')}"
                    if risk.get("impact"):
                        risk_text += f" (Impact: {risk['impact']})"
                    if risk.get("mitigation"):
                        risk_text += f" | Mitigation: {risk['mitigation']}"
                else:
                    risk_text = str(risk)
                doc.add_paragraph(risk_text, style="List Bullet")

        # Management decisions
        decisions = safe_json_parse(update.get("management_decisions", "[]"), [])
        if decisions:
            doc.add_heading(get_string("management_decisions", language), level=3)
            for decision in decisions:
                if isinstance(decision, dict):
                    dec_text = f"{decision.get('decision', '')}"
                    if decision.get("urgency"):
                        dec_text += f" [Urgency: {decision['urgency']}]"
                else:
                    dec_text = str(decision)
                doc.add_paragraph(dec_text, style="List Bullet")

        # Budget status
        if update.get("budget_status") and update.get("budget_status") != "not_applicable":
            doc.add_heading(get_string("budget_status", language), level=3)
            budget_text = update.get("budget_status", "").replace("_", " ").title()
            doc.add_paragraph(budget_text)
            if update.get("budget_notes"):
                doc.add_paragraph(update["budget_notes"], style="Normal")

        # KPIs
        kpis = safe_json_parse(update.get("kpi_updates", "[]"), [])
        if kpis:
            doc.add_heading(get_string("kpi_updates", language), level=3)
            for kpi in kpis:
                if isinstance(kpi, dict):
                    kpi_text = f"{kpi.get('metric', '')}: {kpi.get('value', '')}"
                    if kpi.get("trend"):
                        kpi_text += f" ({kpi['trend']})"
                else:
                    kpi_text = str(kpi)
                doc.add_paragraph(kpi_text, style="List Bullet")

        doc.add_paragraph()  # Spacing between weeks

    doc.add_page_break()

    # === SECTION 3: KEY MILESTONES & ACHIEVEMENTS ===
    heading3 = doc.add_heading(
        get_string("project_goal", language) if language == "de" else "Key Milestones & Achievements",
        level=1
    )
    heading3.runs[0].font.color.rgb = RGBColor(*dark_rgb)

    all_milestones = []
    for update in all_updates:
        milestones = safe_json_parse(update.get("milestone_hit", "[]"), [])
        for milestone in milestones:
            if milestone and milestone not in all_milestones:
                all_milestones.append(milestone)

    if all_milestones:
        for milestone in all_milestones:
            doc.add_paragraph(milestone, style="List Bullet")
    else:
        doc.add_paragraph(get_string("no_data", language))

    doc.add_page_break()

    # === SECTION 4: DECISIONS LOG ===
    heading4 = doc.add_heading(get_string("management_decisions", language), level=1)
    heading4.runs[0].font.color.rgb = RGBColor(*dark_rgb)

    all_decisions = []
    for update in all_updates:
        decisions = safe_json_parse(update.get("management_decisions", "[]"), [])
        week_label = update.get("week_label", "Week")
        for decision in decisions:
            if isinstance(decision, dict):
                all_decisions.append({
                    "week": week_label,
                    "decision": decision,
                })

    if all_decisions:
        for item in all_decisions:
            dec = item["decision"]
            para = doc.add_paragraph(style="List Bullet")
            para.add_run(f"[{item['week']}] ").bold = True
            para.add_run(dec.get("decision", ""))
            if dec.get("urgency"):
                para.add_run(f" ({dec['urgency']})")
    else:
        doc.add_paragraph(get_string("no_data", language))

    doc.add_page_break()

    # === SECTION 5: RISKS & ISSUES LOG ===
    heading5 = doc.add_heading(get_string("risks_blockers", language), level=1)
    heading5.runs[0].font.color.rgb = RGBColor(*dark_rgb)

    all_risks = []
    for update in all_updates:
        risks = safe_json_parse(update.get("risks_blockers", "[]"), [])
        week_label = update.get("week_label", "Week")
        for risk in risks:
            all_risks.append({
                "week": week_label,
                "risk": risk,
            })

    if all_risks:
        for item in all_risks:
            risk = item["risk"]
            para = doc.add_paragraph(style="List Bullet")
            para.add_run(f"[{item['week']}] ").bold = True
            if isinstance(risk, dict):
                para.add_run(f"{risk.get('issue', '')} | Impact: {risk.get('impact', '')} | Mitigation: {risk.get('mitigation', '')}")
            else:
                para.add_run(str(risk))
    else:
        doc.add_paragraph(get_string("no_data", language))

    doc.add_page_break()

    # === SECTION 6: FINAL SUMMARY ===
    heading6 = doc.add_heading(get_string("executive_summary", language), level=1)
    heading6.runs[0].font.color.rgb = RGBColor(*dark_rgb)

    if ai_closure_summary:
        doc.add_paragraph(ai_closure_summary)
    else:
        doc.add_paragraph("No closure summary available.")

    # Save document
    doc.save(filepath)
    return filepath
